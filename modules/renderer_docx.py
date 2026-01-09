# modules/renderer_docx.py
"""
DOCX Renderer - Render documents using docxtpl.
"""
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

from .plugin_loader import PluginPack
from .context_builder import ContextBuilder
from .rule_engine import RuleEngine, EvaluationTrace


# Output directory
OUTPUT_DIR = Path(__file__).parent.parent / "output"


def ensure_output_dir() -> Path:
    """Ensure the output directory exists."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


def hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


class DocxRenderer:
    """Renders DOCX documents from templates and context."""

    def __init__(self, plugin: PluginPack):
        self.plugin = plugin
        self.template_path = plugin.get_template_path()
        self.context_builder = ContextBuilder(plugin)
        self.rule_engine = RuleEngine(plugin)
        self.formatting = plugin.formatting

    def render(
        self,
        data: dict,
        output_path: Optional[Path] = None,
        apply_cell_colors: bool = True
    ) -> tuple[Path, list[EvaluationTrace]]:
        """
        Render a document from the template.

        Args:
            data: Input data dictionary.
            output_path: Optional custom output path.
            apply_cell_colors: Whether to apply cell background colors.

        Returns:
            tuple: (output_path, evaluation_traces)
        """
        # Load template first (needed for Subdoc creation in context building)
        doc = DocxTemplate(str(self.template_path))

        # Build context with template for Subdoc support
        context = self.context_builder.build_context(data, doc)

        # Evaluate rules
        visibility, traces = self.rule_engine.evaluate_all_rules(data)

        # Filter context based on visibility
        context["_visibility"] = visibility
        context["_show_master_file"] = data.get("master_file") == 1

        # Get enabled services
        context["enabled_services"] = self.rule_engine.get_enabled_services(data)

        # Render template with context
        doc.render(context)

        # Apply cell colors if requested
        if apply_cell_colors:
            self._apply_cell_colors(doc.docx, data)

        # Handle [PAGE_BREAK] markers
        self._process_page_breaks(doc.docx)

        # Remove empty paragraphs left by conditional rendering
        self._remove_empty_paragraphs(doc.docx)

        # Remove blank pages
        self._remove_blank_pages(doc.docx)

        # Mark document to update fields (including TOC) when opened
        self._set_update_fields_on_open(doc.docx)

        # Determine output path
        if output_path is None:
            output_path = self._generate_output_path(data)

        # Save document
        ensure_output_dir()
        doc.save(str(output_path))

        return output_path, traces

    def _generate_output_path(self, data: dict) -> Path:
        """Generate output file path based on config pattern."""
        config = self.plugin.config
        pattern = config.get("output", {}).get("filename_pattern", "output_{{ plugin_id }}.docx")

        # Simple pattern replacement
        filename = pattern
        filename = filename.replace("{{ entidad_cliente }}", str(data.get("entidad_cliente", "unknown")))
        filename = filename.replace("{{ anyo_ejercicio }}", str(data.get("anyo_ejercicio", "")))

        # Clean filename
        filename = "".join(c if c.isalnum() or c in "._- " else "_" for c in filename)

        # Add timestamp to ensure uniqueness
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base, ext = os.path.splitext(filename)
        filename = f"{base}_{timestamp}{ext}"

        return OUTPUT_DIR / filename

    def _apply_cell_colors(self, doc: Document, data: dict) -> None:
        """Apply background colors to table cells based on enum values."""
        enum_colors = self.formatting.get("enum_colors", {})

        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip().lower()
                    # Get column context from header row if available
                    column_context = self._get_column_context(table, cell_idx)
                    self._apply_color_if_match(cell, text, enum_colors, column_context)

    def _get_column_context(self, table, cell_idx: int) -> str:
        """Get context from column header to determine color rules."""
        try:
            if len(table.rows) > 0 and cell_idx < len(table.rows[0].cells):
                header_text = table.rows[0].cells[cell_idx].text.strip().lower()
                # Check if this column is for impacto (should not be colored)
                if "impacto" in header_text:
                    return "impacto"
                # Check for compliance columns
                if "cumplimiento" in header_text or "cumplido" in header_text:
                    return "cumplido"
                # Check for afectacion columns
                if "afectación" in header_text or "afectacion" in header_text:
                    return "afectacion"
        except Exception:
            pass
        return ""

    def _apply_color_if_match(self, cell, text: str, enum_colors: dict, column_context: str = "") -> None:
        """Apply color to a cell if the text matches an enum value."""
        # Skip coloring for impacto columns
        if column_context == "impacto":
            return

        # Normalize text: treat checkmark as "si" for color matching
        normalized_text = text
        if text == "✓":
            normalized_text = "si"

        # Check afectacion_color first (most specific - bajo/medio/alto)
        afectacion_map = enum_colors.get("afectacion_color", {}).get("values", {})
        if normalized_text in afectacion_map:
            color_def = afectacion_map[normalized_text]
            if "background" in color_def:
                self._set_cell_background(cell, color_def.get("background"))
            return

        # Check cumplido_color (si/no/parcial)
        cumplido_map = enum_colors.get("cumplido_color", {}).get("values", {})
        if normalized_text in cumplido_map:
            color_def = cumplido_map[normalized_text]
            if "background" in color_def:
                self._set_cell_background(cell, color_def.get("background"))
            return

        # Check compliance_color (si/no for summary tables)
        compliance_map = enum_colors.get("compliance_color", {}).get("values", {})
        if normalized_text in compliance_map:
            color_def = compliance_map[normalized_text]
            if "background" in color_def:
                self._set_cell_background(cell, color_def.get("background"))
            return

    def _set_cell_background(self, cell, hex_color: str) -> None:
        """Set the background color of a table cell."""
        try:
            # Remove # if present
            hex_color = hex_color.lstrip("#")

            # Create shading element
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
            )

            # Apply to cell
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception:
            pass  # Silently fail if color cannot be applied

    def _process_page_breaks(self, doc: Document) -> None:
        """
        Process [PAGE_BREAK] markers in the document.
        Replace them with actual page breaks and remove the marker text.
        """
        PAGE_BREAK_MARKER = "[PAGE_BREAK]"

        for paragraph in doc.paragraphs:
            if PAGE_BREAK_MARKER in paragraph.text:
                # Find the run containing the marker
                for run in paragraph.runs:
                    if PAGE_BREAK_MARKER in run.text:
                        # Replace marker with page break
                        parts = run.text.split(PAGE_BREAK_MARKER)
                        run.text = parts[0]

                        # Add page break
                        run.add_break(WD_BREAK.PAGE)

                        # Add remaining text if any
                        if len(parts) > 1 and parts[1].strip():
                            run.add_text(parts[1])
                        break

        # Also check in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if PAGE_BREAK_MARKER in paragraph.text:
                            for run in paragraph.runs:
                                if PAGE_BREAK_MARKER in run.text:
                                    parts = run.text.split(PAGE_BREAK_MARKER)
                                    run.text = parts[0]
                                    run.add_break(WD_BREAK.PAGE)
                                    if len(parts) > 1 and parts[1].strip():
                                        run.add_text(parts[1])
                                    break

    def _set_update_fields_on_open(self, doc: Document) -> None:
        """
        Set the document to update all fields (including TOC) when opened in Word.

        This adds the w:updateFields setting to the document properties,
        which prompts Word to update the Table of Contents and other fields
        when the document is first opened.
        """
        try:
            # Access the document's settings element
            settings = doc.settings.element

            # Create or find the updateFields element
            # Namespace for Word settings
            w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

            # Check if updateFields element exists
            update_fields = settings.find(f"{w_ns}updateFields")

            if update_fields is None:
                # Create the updateFields element
                update_fields = OxmlElement("w:updateFields")
                update_fields.set(f"{w_ns}val", "true")
                settings.append(update_fields)
            else:
                # Update existing element
                update_fields.set(f"{w_ns}val", "true")
        except Exception:
            # Silently fail if we can't set this property
            pass

    def _remove_empty_paragraphs(self, doc: Document) -> None:
        """
        Remove empty paragraphs left by conditional rendering.

        When {% if condition %} blocks don't match, they may leave behind
        empty paragraphs. This method removes:
        1. Completely empty paragraphs (no text, no images)
        2. Consecutive empty paragraphs (keeps only one if between content)
        """
        try:
            body = doc.element.body
            w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

            # Get all paragraph elements directly under body
            paragraphs = list(body.findall(f"{w_ns}p"))

            def is_paragraph_empty(p_elem):
                """Check if a paragraph is completely empty."""
                text_content = "".join(p_elem.itertext()).strip()
                if text_content:
                    return False

                # Check for meaningful content (images, objects, etc.)
                for child in p_elem.iter():
                    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                    if tag in ("drawing", "pict", "object", "br"):
                        return False

                return True

            def has_section_break(p_elem):
                """Check if paragraph has section properties (section break)."""
                return p_elem.find(f".//{w_ns}sectPr") is not None

            # Track elements to remove
            elements_to_remove = []

            # Find consecutive empty paragraphs
            prev_was_empty = False
            for i, p in enumerate(paragraphs):
                if is_paragraph_empty(p):
                    # Don't remove paragraphs with section breaks
                    if has_section_break(p):
                        prev_was_empty = False
                        continue

                    # If previous paragraph was also empty, mark this one for removal
                    if prev_was_empty:
                        elements_to_remove.append(p)
                    # If this is a standalone empty paragraph not at the start/end
                    elif i > 0 and i < len(paragraphs) - 1:
                        # Check if next paragraph is also empty
                        next_p = paragraphs[i + 1]
                        if is_paragraph_empty(next_p):
                            # This is part of a sequence - keep the first one
                            pass
                        else:
                            # Standalone empty paragraph between content
                            # Check if previous had content - if so, might be intentional spacing
                            prev_p = paragraphs[i - 1]
                            if is_paragraph_empty(prev_p):
                                # Previous was also empty, this is consecutive
                                elements_to_remove.append(p)

                    prev_was_empty = True
                else:
                    prev_was_empty = False

            # Also check for empty paragraphs at the very beginning or end
            # that serve no purpose
            if paragraphs:
                # Check first paragraph
                if is_paragraph_empty(paragraphs[0]) and not has_section_break(paragraphs[0]):
                    if len(paragraphs) > 1 and is_paragraph_empty(paragraphs[1]):
                        elements_to_remove.append(paragraphs[0])

            # Remove identified elements
            for elem in elements_to_remove:
                parent = elem.getparent()
                if parent is not None:
                    try:
                        parent.remove(elem)
                    except Exception:
                        pass

        except Exception:
            # Silently fail if empty paragraph removal encounters issues
            pass

    def _remove_blank_pages(self, doc: Document) -> None:
        """
        Remove blank pages from the document.

        A blank page is detected by:
        1. Empty paragraphs between page breaks
        2. Sections with only whitespace content
        3. Consecutive page breaks
        """
        try:
            body = doc.element.body
            w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

            # Track elements to remove
            elements_to_remove = []

            # Get all paragraph elements
            paragraphs = body.findall(f".//{w_ns}p")

            # Helper to check if a paragraph is essentially empty
            def is_paragraph_empty(p_elem):
                """Check if a paragraph contains only whitespace or is empty."""
                text_content = "".join(p_elem.itertext()).strip()
                if text_content:
                    return False

                # Check if it has any meaningful content (images, tables, etc.)
                # but allow page breaks
                for child in p_elem.iter():
                    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                    if tag in ("drawing", "pict", "object"):
                        return False
                return True

            # Helper to check if paragraph has a page break
            def has_page_break(p_elem):
                """Check if paragraph contains a page break."""
                for br in p_elem.findall(f".//{w_ns}br"):
                    br_type = br.get(f"{w_ns}type")
                    if br_type == "page":
                        return True
                return False

            # Find sequences of empty paragraphs followed by page breaks
            i = 0
            while i < len(paragraphs):
                p = paragraphs[i]

                # Check for pattern: empty paragraph with page break
                if is_paragraph_empty(p) and has_page_break(p):
                    # Look ahead to see if next content is also empty
                    j = i + 1
                    consecutive_empty = True

                    # Check next few paragraphs for content
                    while j < len(paragraphs) and j < i + 5:
                        next_p = paragraphs[j]
                        if not is_paragraph_empty(next_p):
                            consecutive_empty = False
                            break
                        if has_page_break(next_p):
                            # Found another page break after empty content - remove the break
                            elements_to_remove.append(p)
                            break
                        j += 1

                i += 1

            # Also find and remove empty section breaks that would create blank pages
            sect_prs = body.findall(f".//{w_ns}sectPr")
            for sect_pr in sect_prs:
                parent = sect_pr.getparent()
                if parent is not None:
                    # Check if the section is in an empty paragraph
                    if parent.tag.endswith("}p"):
                        text_content = "".join(parent.itertext()).strip()
                        if not text_content:
                            # Check if previous content is also empty
                            prev = parent.getprevious()
                            if prev is not None:
                                prev_text = "".join(prev.itertext()).strip()
                                if not prev_text and prev not in elements_to_remove:
                                    elements_to_remove.append(prev)

            # Remove identified elements
            for elem in elements_to_remove:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)

        except Exception:
            # Silently fail if blank page removal encounters issues
            pass


def render_document(
    plugin: PluginPack,
    data: dict,
    output_path: Optional[Path] = None
) -> tuple[Path, list[EvaluationTrace]]:
    """
    Convenience function to render a document.

    Args:
        plugin: Loaded plugin pack.
        data: Input data dictionary.
        output_path: Optional custom output path.

    Returns:
        tuple: (output_path, evaluation_traces)
    """
    renderer = DocxRenderer(plugin)
    return renderer.render(data, output_path)
