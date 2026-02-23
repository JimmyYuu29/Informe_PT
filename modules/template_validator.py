# modules/template_validator.py
"""
Template Validator - Validates uploaded DOCX templates for correctness.

Performs the following checks:
  A. Jinja2/docxtpl syntax parsability
  B. Variable consistency (template vars vs plugin available fields)
  C. Smoke-test rendering with sample data
  D. Anchor/keyword structure protection (optional)
"""
import hashlib
import json
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

from docxtpl import DocxTemplate
from docx import Document


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class ValidationIssue:
    """A single validation issue."""
    level: str  # "error", "warning", "info"
    category: str  # "syntax", "variables", "render", "anchor"
    message: str
    suggestion: Optional[str] = None

    def to_dict(self) -> dict:
        d = {"level": self.level, "category": self.category, "message": self.message}
        if self.suggestion:
            d["suggestion"] = self.suggestion
        return d


@dataclass
class TemplateValidationResult:
    """Result of template validation."""
    status: str = "PASS"  # PASS, WARN, FAIL
    issues: list[ValidationIssue] = field(default_factory=list)
    variables_found: list[str] = field(default_factory=list)
    variables_missing: list[str] = field(default_factory=list)
    variables_extra: list[str] = field(default_factory=list)
    sha256: str = ""
    timestamp: str = ""

    def add_error(self, category: str, message: str, suggestion: str = None) -> None:
        self.issues.append(ValidationIssue("error", category, message, suggestion))
        self.status = "FAIL"

    def add_warning(self, category: str, message: str, suggestion: str = None) -> None:
        self.issues.append(ValidationIssue("warning", category, message, suggestion))
        if self.status == "PASS":
            self.status = "WARN"

    def add_info(self, category: str, message: str) -> None:
        self.issues.append(ValidationIssue("info", category, message))

    @property
    def errors(self) -> list[ValidationIssue]:
        return [i for i in self.issues if i.level == "error"]

    @property
    def warnings(self) -> list[ValidationIssue]:
        return [i for i in self.issues if i.level == "warning"]

    def to_dict(self) -> dict:
        return {
            "status": self.status,
            "issues": [i.to_dict() for i in self.issues],
            "variables_found": self.variables_found,
            "variables_missing": self.variables_missing,
            "variables_extra": self.variables_extra,
            "sha256": self.sha256,
            "timestamp": self.timestamp,
            "errors_count": len(self.errors),
            "warnings_count": len(self.warnings),
        }


# ============================================================================
# Helper Functions
# ============================================================================

def compute_sha256(file_bytes: bytes) -> str:
    """Compute SHA-256 hash of file bytes."""
    return hashlib.sha256(file_bytes).hexdigest()


def extract_template_variables(template_path: str) -> set[str]:
    """
    Extract Jinja2 variable names from a DOCX template.

    Finds {{ variable }}, {% if variable %}, {% for x in variable %}, etc.

    Returns:
        Set of variable names found in the template.
    """
    variables = set()

    try:
        doc = Document(template_path)

        # Collect all text from the document
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)

        full_text = "\n".join(all_text)

        # Extract {{ variable }} patterns
        # Match simple variables: {{ var_name }}
        var_pattern = re.compile(r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_.]*(?:\s*\|[^}]*)?)\s*\}\}')
        for match in var_pattern.finditer(full_text):
            var_name = match.group(1).split("|")[0].strip()
            # Skip loop variables and special constructs
            if "." in var_name:
                # e.g., item.field - extract the base variable
                base = var_name.split(".")[0]
                variables.add(base)
            else:
                variables.add(var_name)

        # Extract {% if variable %} patterns
        if_pattern = re.compile(r'\{%[-\s]*if\s+(.+?)\s*[-]?%\}')
        for match in if_pattern.finditer(full_text):
            condition = match.group(1)
            # Extract variable names from conditions
            cond_vars = re.findall(r'([a-zA-Z_][a-zA-Z0-9_]*)', condition)
            for v in cond_vars:
                if v not in ('and', 'or', 'not', 'in', 'is', 'true', 'false',
                             'none', 'True', 'False', 'None', 'if', 'else',
                             'elif', 'endif', 'for', 'endfor', 'set'):
                    variables.add(v)

        # Extract {% for x in variable %} patterns
        for_pattern = re.compile(r'\{%[-\s]*(?:tr\s+)?for\s+(\w+)\s+in\s+(\w+)')
        for match in for_pattern.finditer(full_text):
            loop_var = match.group(1)
            iterable = match.group(2)
            variables.add(iterable)
            # Remove loop variable from variables as it's locally scoped
            variables.discard(loop_var)

        # Extract {%tr for x in variable %} patterns (table row loops)
        tr_for_pattern = re.compile(r'\{%tr\s+for\s+(\w+)\s+in\s+(\w+)')
        for match in tr_for_pattern.finditer(full_text):
            loop_var = match.group(1)
            iterable = match.group(2)
            variables.add(iterable)
            variables.discard(loop_var)

    except Exception:
        pass

    # Remove common Jinja2 built-ins and internal variables
    builtins = {
        'loop', 'range', 'true', 'false', 'none', 'True', 'False', 'None',
        '_visibility', '_show_master_file', 'enabled_services',
    }
    variables -= builtins

    return variables


def get_available_fields(plugin) -> set[str]:
    """
    Get all available fields from a plugin pack.

    Combines:
    - Field definitions from fields.yaml
    - Derived fields from derived.yaml
    - Text block keys from texts.yaml
    - Table keys from tables.yaml
    - Special context variables added by context_builder

    Args:
        plugin: PluginPack instance.

    Returns:
        Set of all available field names.
    """
    available = set()

    # Fields from fields.yaml
    fields_def = plugin.get_field_definitions()
    available.update(fields_def.keys())

    # Derived fields
    derived_def = plugin.get_derived_fields()
    available.update(derived_def.keys())

    # Text block keys
    text_blocks = plugin.get_text_blocks()
    available.update(text_blocks.keys())

    # Table keys
    tables = plugin.get_table_definitions()
    available.update(tables.keys())

    # Common context variables added by the system
    system_vars = {
        # Standard context variables
        '_visibility', '_show_master_file', 'enabled_services',
        'plugin_id', 'plugin_name', 'generation_date',
        'anyo_ejercicio', 'anyo_anterior',
        # Financial derived
        'tabla_num',
        # Comentarios
        'comentarios_valorativos',
        # OOVV
        'servicios_vinculados', 'servicios_oovv',
        'documentacion_facilitada',
        'total_ingresos', 'total_gastos',
        'peso_oovv_incn', 'peso_oovv_costes',
        'valoracion_oovv',
        # Contacts
        'contacto1', 'contacto2', 'contacto3',
        'cargo_contacto1', 'cargo_contacto2', 'cargo_contacto3',
        'correo_contacto1', 'correo_contacto2', 'correo_contacto3',
    }
    available.update(system_vars)

    # Risk table variables (impacto_1..12, afectacion_pre_1..12, etc.)
    for i in range(1, 13):
        available.update([
            f"impacto_{i}", f"afectacion_pre_{i}",
            f"texto_mitigacion_{i}", f"afectacion_final_{i}",
        ])

    # Compliance variables
    for prefix in ("local", "mast"):
        for i in range(1, 18):
            available.update([
                f"cumplido_{prefix}_{i}",
                f"texto_cumplido_{prefix}_{i}",
            ])
        for i in range(1, 5):
            available.add(f"cumplimiento_resumen_{prefix}_{i}")

    return available


# ============================================================================
# Validation Steps
# ============================================================================

def validate_syntax(template_path: str, result: TemplateValidationResult) -> bool:
    """
    Step A: Validate Jinja2/docxtpl syntax by attempting to load the template.

    Returns True if syntax is parsable.
    """
    try:
        doc = DocxTemplate(template_path)
        result.add_info("syntax", "Template loaded and parsed successfully (docxtpl)")
        return True
    except Exception as e:
        result.add_error(
            "syntax",
            f"Template syntax error: {e}",
            "Check Jinja2 syntax in the template. Ensure all {{ }}, {% %} blocks are properly closed."
        )
        return False


def validate_variables(template_path: str, plugin, result: TemplateValidationResult) -> None:
    """
    Step B: Validate variable consistency between template and plugin fields.
    """
    template_vars = extract_template_variables(template_path)
    available_fields = get_available_fields(plugin)

    result.variables_found = sorted(template_vars)

    # Variables in template but not in plugin (potentially unknown)
    extra_vars = template_vars - available_fields
    result.variables_extra = sorted(extra_vars)

    # Variables in plugin but not in template (potentially missing from template)
    # This is informational, not an error
    missing_vars = available_fields - template_vars

    if extra_vars:
        for var in sorted(extra_vars):
            result.add_warning(
                "variables",
                f"Template variable '{var}' not found in plugin field definitions",
                f"Ensure '{var}' is a valid context variable or derived field."
            )

    result.add_info(
        "variables",
        f"Found {len(template_vars)} template variables, "
        f"{len(available_fields)} available fields, "
        f"{len(extra_vars)} unmatched"
    )


def validate_render_smoke(
    template_path: str,
    plugin,
    sample_input_path: Optional[str],
    result: TemplateValidationResult,
) -> None:
    """
    Step C: Smoke-test rendering with sample data.
    Checks:
    - Rendering doesn't raise exceptions
    - Output DOCX doesn't contain residual {{ or {% markers
    """
    try:
        # Load sample data
        sample_data = {}
        if sample_input_path and Path(sample_input_path).exists():
            with open(sample_input_path, "r", encoding="utf-8") as f:
                sample_data = json.load(f)
        else:
            # Generate minimal sample data from plugin fields
            fields_def = plugin.get_field_definitions()
            for fname, fdef in fields_def.items():
                ftype = fdef.get("type", "text")
                if ftype == "text":
                    sample_data[fname] = f"sample_{fname}"
                elif ftype == "int":
                    sample_data[fname] = 0
                elif ftype in ("decimal", "currency", "percentage"):
                    sample_data[fname] = 0.0
                elif ftype == "bool":
                    sample_data[fname] = False
                elif ftype == "enum":
                    values = fdef.get("values", [])
                    if values:
                        if isinstance(values[0], dict):
                            sample_data[fname] = values[0].get("value", "")
                        else:
                            sample_data[fname] = values[0]
                elif ftype == "list":
                    sample_data[fname] = []
                elif ftype == "date":
                    sample_data[fname] = "2025-01-01"

        # Attempt to render
        doc = DocxTemplate(template_path)

        # Build a basic context (simplified - we don't use full context_builder
        # to avoid side effects)
        from .context_builder import ContextBuilder
        from .generate import preprocess_input

        processed = preprocess_input(sample_data)
        context_builder = ContextBuilder(plugin)
        context = context_builder.build_context(processed, doc)

        # Add visibility defaults
        context.setdefault("_visibility", {})
        context.setdefault("_show_master_file", False)
        context.setdefault("enabled_services", [])

        doc.render(context)

        # Save to temp file and check for residual markers
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)

        # Check for residual {{ or {% in output
        out_doc = Document(tmp_path)
        residual_found = False
        residual_markers = []

        for para in out_doc.paragraphs:
            text = para.text
            if "{{" in text or "{%" in text:
                residual_found = True
                residual_markers.append(text[:100])

        for table in out_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        if "{{" in text or "{%" in text:
                            residual_found = True
                            residual_markers.append(text[:100])

        # Clean up temp file
        try:
            Path(tmp_path).unlink()
        except Exception:
            pass

        if residual_found:
            for marker in residual_markers[:5]:
                result.add_warning(
                    "render",
                    f"Residual template marker found in output: '{marker}'",
                    "This variable may not be populated during rendering. "
                    "Check the template variable name or context builder."
                )
        else:
            result.add_info("render", "Smoke-test rendering completed successfully with no residual markers")

    except Exception as e:
        result.add_error(
            "render",
            f"Smoke-test rendering failed: {e}",
            "Ensure all template variables are provided in the context. "
            "Check for undefined variables or incorrect loop syntax."
        )


def validate_anchors(
    template_path: str,
    result: TemplateValidationResult,
    required_anchors: Optional[list[str]] = None,
    strict: bool = False,
) -> None:
    """
    Step D: Anchor/keyword structure protection.
    Checks that certain keywords/anchors exist in the template.
    """
    if not required_anchors:
        # Default anchors for PT Review template
        required_anchors = [
            "Contactos",
        ]

    try:
        doc = Document(template_path)
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)

        full_text = "\n".join(all_text)

        for anchor in required_anchors:
            if anchor.lower() not in full_text.lower():
                if strict:
                    result.add_error(
                        "anchor",
                        f"Required anchor/keyword '{anchor}' not found in template",
                        f"The template must contain the keyword '{anchor}' for document structure integrity."
                    )
                else:
                    result.add_warning(
                        "anchor",
                        f"Expected anchor/keyword '{anchor}' not found in template",
                        f"Consider including '{anchor}' for document structure consistency."
                    )
            else:
                result.add_info("anchor", f"Anchor '{anchor}' found in template")

    except Exception as e:
        result.add_warning(
            "anchor",
            f"Could not check anchors: {e}",
        )


# ============================================================================
# Main Validation Entry Point
# ============================================================================

def validate_template(
    template_path: str,
    plugin,
    sample_input_path: Optional[str] = None,
    check_anchors: bool = True,
    required_anchors: Optional[list[str]] = None,
    strict_anchors: bool = False,
) -> TemplateValidationResult:
    """
    Full template validation pipeline.

    Args:
        template_path: Path to the DOCX template file.
        plugin: PluginPack instance.
        sample_input_path: Optional path to sample input JSON for smoke test.
        check_anchors: Whether to check for required anchors.
        required_anchors: List of required anchor keywords.
        strict_anchors: If True, missing anchors are errors; if False, warnings.

    Returns:
        TemplateValidationResult with all validation details.
    """
    result = TemplateValidationResult()
    result.timestamp = datetime.utcnow().isoformat() + "Z"

    # Compute file hash
    try:
        with open(template_path, "rb") as f:
            file_bytes = f.read()
        result.sha256 = compute_sha256(file_bytes)
    except Exception as e:
        result.add_error("general", f"Cannot read template file: {e}")
        return result

    # Step A: Syntax validation
    syntax_ok = validate_syntax(template_path, result)
    if not syntax_ok:
        return result

    # Step B: Variable consistency
    validate_variables(template_path, plugin, result)

    # Step C: Smoke-test rendering
    if sample_input_path is None:
        # Try to find default sample input
        project_root = Path(__file__).parent.parent
        default_sample = project_root / "tests" / "golden" / "sample_input.json"
        if default_sample.exists():
            sample_input_path = str(default_sample)

    validate_render_smoke(template_path, plugin, sample_input_path, result)

    # Step D: Anchor protection
    if check_anchors:
        validate_anchors(template_path, result, required_anchors, strict_anchors)

    return result


def generate_validation_report(result: TemplateValidationResult) -> dict:
    """Generate a validation report as a dictionary (for JSON serialization)."""
    return result.to_dict()
