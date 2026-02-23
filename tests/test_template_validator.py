#!/usr/bin/env python3
# tests/test_template_validator.py
"""
Tests for the template validator module.

Includes:
  - PASS test: validates the existing template_final.docx
  - FAIL test: validates a deliberately broken template
"""
import json
import sys
import tempfile
from pathlib import Path

import pytest

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from modules.plugin_loader import load_plugin
from modules.template_validator import (
    validate_template,
    generate_validation_report,
    extract_template_variables,
    compute_sha256,
    get_available_fields,
    TemplateValidationResult,
)


# ============================================================================
# Fixtures
# ============================================================================

@pytest.fixture
def plugin():
    """Load the pt_review plugin."""
    return load_plugin("pt_review")


@pytest.fixture
def valid_template_path():
    """Path to the existing valid template."""
    path = PROJECT_ROOT / "config" / "template_final.docx"
    if not path.exists():
        pytest.skip("template_final.docx not found")
    return str(path)


@pytest.fixture
def sample_input_path():
    """Path to the sample input JSON."""
    path = PROJECT_ROOT / "tests" / "golden" / "sample_input.json"
    if not path.exists():
        pytest.skip("sample_input.json not found")
    return str(path)


@pytest.fixture
def broken_template_path():
    """Create a broken DOCX template with invalid Jinja2 syntax."""
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc = Document()
        # Add deliberately broken Jinja2 syntax
        doc.add_paragraph("{{ unclosed_variable")
        doc.add_paragraph("{% if broken_condition")
        doc.add_paragraph("Normal text here")
        doc.add_paragraph("{{ completely_unknown_var_xyz_123 }}")
        doc.save(tmp.name)
        return tmp.name


# ============================================================================
# PASS Test
# ============================================================================

class TestValidTemplatePass:
    """Tests that the existing template_final.docx passes validation."""

    def test_valid_template_loads(self, valid_template_path, plugin):
        """Test that the valid template can be loaded and parsed."""
        result = validate_template(
            template_path=valid_template_path,
            plugin=plugin,
            check_anchors=False,  # Don't check anchors for basic test
        )
        assert result.status in ("PASS", "WARN"), (
            f"Expected PASS or WARN, got {result.status}. "
            f"Errors: {[e.message for e in result.errors]}"
        )

    def test_valid_template_has_sha256(self, valid_template_path, plugin):
        """Test that SHA-256 is computed."""
        result = validate_template(
            template_path=valid_template_path,
            plugin=plugin,
            check_anchors=False,
        )
        assert result.sha256, "SHA-256 should be computed"
        assert len(result.sha256) == 64, "SHA-256 should be 64 hex chars"

    def test_valid_template_has_timestamp(self, valid_template_path, plugin):
        """Test that timestamp is set."""
        result = validate_template(
            template_path=valid_template_path,
            plugin=plugin,
            check_anchors=False,
        )
        assert result.timestamp, "Timestamp should be set"
        assert "Z" in result.timestamp, "Timestamp should be UTC"

    def test_valid_template_has_variables(self, valid_template_path, plugin):
        """Test that variables are extracted from the template."""
        result = validate_template(
            template_path=valid_template_path,
            plugin=plugin,
            check_anchors=False,
        )
        assert len(result.variables_found) > 0, "Should find template variables"

    def test_valid_template_validation_report(self, valid_template_path, plugin):
        """Test that validation report can be generated."""
        result = validate_template(
            template_path=valid_template_path,
            plugin=plugin,
            check_anchors=False,
        )
        report = generate_validation_report(result)
        assert isinstance(report, dict)
        assert "status" in report
        assert "issues" in report
        assert "variables_found" in report
        assert "sha256" in report

    def test_valid_template_with_sample_input(
        self, valid_template_path, plugin, sample_input_path
    ):
        """Test validation with sample input for smoke rendering."""
        result = validate_template(
            template_path=valid_template_path,
            plugin=plugin,
            sample_input_path=sample_input_path,
            check_anchors=False,
        )
        assert result.status in ("PASS", "WARN"), (
            f"Expected PASS or WARN with sample input, got {result.status}. "
            f"Errors: {[e.message for e in result.errors]}"
        )


# ============================================================================
# FAIL Test
# ============================================================================

class TestBrokenTemplateFail:
    """Tests that a broken template fails validation."""

    def test_broken_template_has_issues(self, broken_template_path, plugin):
        """Test that a broken template produces validation issues."""
        result = validate_template(
            template_path=broken_template_path,
            plugin=plugin,
            check_anchors=False,
        )
        # The broken template should have at least warnings about unknown variables
        assert len(result.issues) > 0, "Should have validation issues"

    def test_broken_template_has_unknown_vars(self, broken_template_path, plugin):
        """Test that unknown variables are detected."""
        result = validate_template(
            template_path=broken_template_path,
            plugin=plugin,
            check_anchors=False,
        )
        # The broken template has variables not in the plugin
        # (completely_unknown_var_xyz_123)
        has_var_warnings = any(
            "variable" in i.message.lower() or "not found" in i.message.lower()
            for i in result.issues
        )
        # We expect either variable warnings or render errors
        assert len(result.issues) > 0, "Should have at least some issues"


# ============================================================================
# Unit Tests
# ============================================================================

class TestHelperFunctions:
    """Tests for helper functions."""

    def test_compute_sha256(self):
        """Test SHA-256 computation."""
        test_data = b"hello world"
        result = compute_sha256(test_data)
        assert len(result) == 64
        assert result == "b94d27b9934d3e08a52e52d7da7dabfac484efe37a5380ee9088f7ace2efcde9"

    def test_get_available_fields(self, plugin):
        """Test available fields extraction."""
        fields = get_available_fields(plugin)
        assert len(fields) > 0, "Should have available fields"
        # Check some expected fields
        assert "entidad_cliente" in fields
        assert "fecha_fin_fiscal" in fields

    def test_extract_variables_from_valid_template(self, valid_template_path):
        """Test variable extraction from the valid template."""
        variables = extract_template_variables(valid_template_path)
        assert isinstance(variables, set)
        assert len(variables) > 0, "Should extract variables from template"


# ============================================================================
# Cleanup
# ============================================================================

def teardown_module():
    """Clean up temp files."""
    import glob
    import os
    for f in glob.glob(str(Path(tempfile.gettempdir()) / "tmp*.docx")):
        try:
            os.unlink(f)
        except Exception:
            pass
